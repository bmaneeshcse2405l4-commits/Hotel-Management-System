from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def main():
    doc = Document()
    
    # Title
    heading = doc.add_heading('Table 2.1: Literature Survey Table', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Define Column Widths (approximation)
    widths = [Inches(0.4), Inches(1.2), Inches(1.8), Inches(0.6), Inches(1.8), Inches(1.8)]
    
    hdr_cells = table.rows[0].cells
    headers = ['S. No', 'Author(s)', 'Title', 'Year', 'Merits', 'Demerits']
    
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ("1", "A. M. Chopde; Manish Arote", "Automated and Efficient Hotel Management System", "2024", 
         "Centralized system simplifies reservations, staff, and guest management.", "Desktop-based; limited remote accessibility."),
        
        ("2", "Ogirima S. A. O.; Awode T. R.; Adeosun O. O.", "Online Computerized Hotel Management System", "2014", 
         "Online reservations improve speed, reliability, and accessibility.", "Case-study limited to a single hotel."),
        
        ("3", "W.P.S.W. Weerasinghe et al.", "Research on Hotel Management System", "2022", 
         "Improved reservation, food ordering, and employee management.", "Requires internet dependency."),
        
        ("4", "Tarunesh Gautam; Satyam Gaurav", "A Research on Hotel Management System", "2022", 
         "Efficient guest booking and information management.", "Basic system with limited scalability."),
        
        ("5", "Sahil; Deepak Kumar; Surender", "A Study on Hotel Management System", "2023", 
         "Automation increases operational efficiency and customer satisfaction.", "Security and cloud issues not deeply explored."),
        
        ("6", "Jeevanjot Singh; Amanjot Singh Mavi; Kiranpreet Kaur", "Comprehensive Review on Web-Based Hotel Management System", "2023", 
         "Web systems enhance revenue and guest experience.", "Implementation complexity."),
        
        ("7", "Nigamananda Pradhan; Rakesh Pradhan", "Hotel Management System", "2025", 
         "Automation improves productivity and customer satisfaction.", "Needs continuous maintenance."),
        
        ("8", "Dussip Eldar", "Creating a Website for Hotel 'Ertis'", "2022", 
         "Online booking reduces manual errors and processing time.", "Limited evaluation scope."),
        
        ("9", "Group-F ICT Department", "Project on Hotel Management System", "2022", 
         "Online booking and admin approval improves flexibility.", "Academic prototype only."),
        
        ("10", "Multiple Authors", "Automated Hotel Information Management Studies", "2024", 
         "Centralized information improves operational control.", "Requires organizational adaptation.")
    ]
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    # Apply column widths manually
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    doc.save('Literature_Survey_Table.docx')
    print("Literature Survey Table Document Successfully Generated!")

if __name__ == '__main__':
    main()
